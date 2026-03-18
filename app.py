import streamlit as st
from docx import Document
from io import BytesIO
from datetime import date, datetime
from pypdf import PdfReader
import os
from openai import OpenAI
import base64
from io import BytesIO
from docx import Document
import base64
from io import BytesIO
import fitz  # PyMuPDF
import json
import re
 
# =============================
# CONFIG INICIAL
# =============================
st.set_page_config(page_title="Sistema Interno - Estudio Peire", layout="wide")

# =============================
# ESTILO VISUAL
# =============================
PRIMARY = "#1597C0"
PRIMARY_DARK = "#0F7EA0"
BG = "#F5F6F7"
CARD = "#FFFFFF"
TEXT = "#1F2937"
MUTED = "#4B5563"
BORDER = "#D6E2E8"

def aplicar_estilo():
    st.markdown(
        f"""
        <style>
        .stApp {{
            background-color: {BG};
            color: {TEXT};
        }}

        [data-testid="stSidebar"] {{
            background-color: #ffffff;
            border-right: 1px solid {BORDER};
        }}

        [data-testid="stSidebar"] * {{
            color: {TEXT} !important;
        }}

        h1, h2, h3 {{
            color: {PRIMARY_DARK};
            font-weight: 700;
        }}

        label, .stMarkdown p, .stMarkdown li, .stMarkdown span {{
            color: {TEXT};
        }}

        pre, code, .stCode, [data-testid="stCodeBlock"], [data-testid="stCodeBlock"] * {{
            background-color: #f8fafc !important;
            color: #111827 !important;
            border-radius: 12px !important;
        }}

        .bloque-header {{
            background: linear-gradient(90deg, #ffffff 0%, #f7fbfd 100%);
            border: 1px solid {BORDER};
            border-radius: 18px;
            padding: 22px 26px;
            margin-bottom: 18px;
            box-shadow: 0 2px 10px rgba(0,0,0,0.04);
        }}

        .titulo-principal {{
            font-size: 2.1rem;
            font-weight: 800;
            color: {PRIMARY_DARK};
            margin-bottom: 0.2rem;
        }}

        .subtitulo-principal {{
            font-size: 1rem;
            color: {MUTED};
            margin-top: 0;
        }}

        .mini-card {{
            background-color: {CARD};
            border: 1px solid {BORDER};
            border-radius: 16px;
            padding: 18px;
            margin-bottom: 14px;
            box-shadow: 0 2px 8px rgba(0,0,0,0.03);
        }}

        .mini-card h4 {{
            color: {PRIMARY_DARK};
            margin-bottom: 8px;
        }}

        .mini-card p {{
            color: {MUTED};
            margin-bottom: 0;
            font-size: 0.95rem;
        }}

        .login-wrap {{
            display: flex;
            justify-content: center;
            margin-top: 30px;
            margin-bottom: 10px;
        }}

        .login-box {{
            width: 100%;
            max-width: 520px;
            background: #ffffff;
            padding: 32px 28px;
            border-radius: 18px;
            border: 1px solid {BORDER};
            box-shadow: 0 4px 20px rgba(0,0,0,0.05);
            text-align: center;
        }}

        .login-title {{
            font-size: 2rem;
            font-weight: 800;
            color: {PRIMARY_DARK};
            margin-bottom: 0.2rem;
        }}

        .login-subtitle {{
            color: {MUTED};
            margin-bottom: 0.2rem;
        }}

        .stButton > button {{
            background-color: {PRIMARY};
            color: white !important;
            border: none !important;
            border-radius: 12px !important;
            padding: 0.65rem 1rem !important;
            font-weight: 600 !important;
            width: 100%;
            opacity: 1 !important;
        }}

        .stButton > button:hover {{
            background-color: {PRIMARY_DARK} !important;
            color: white !important;
        }}

        .stDownloadButton > button {{
            background-color: {PRIMARY} !important;
            color: white !important;
            border: none !important;
            border-radius: 12px !important;
            padding: 0.65rem 1rem !important;
            font-weight: 600 !important;
            width: 100%;
            opacity: 1 !important;
        }}

        .stDownloadButton > button:hover {{
            background-color: {PRIMARY_DARK} !important;
            color: white !important;
        }}

        .stDownloadButton button p,
        .stDownloadButton button span,
        .stDownloadButton button div {{
            color: white !important;
            opacity: 1 !important;
        }}

        .stTextInput input,
        .stTextArea textarea {{
            background-color: #ffffff !important;
            color: {TEXT} !important;
            border: 1px solid {BORDER} !important;
            border-radius: 12px !important;
            box-shadow: none !important;
            caret-color: {TEXT} !important;
        }}

        .stTextInput input:focus,
        .stTextArea textarea:focus {{
            border: 1px solid {PRIMARY} !important;
            box-shadow: 0 0 0 1px {PRIMARY} !important;
            outline: none !important;
        }}

        .stTextInput input::selection,
        .stTextArea textarea::selection {{
            background-color: #bae6fd !important;
            color: {TEXT} !important;
        }}

        .stTextInput input::placeholder,
        .stTextArea textarea::placeholder {{
            color: {MUTED} !important;
            opacity: 1 !important;
        }}

        div[data-baseweb="select"] > div {{
            background-color: #ffffff !important;
            color: {TEXT} !important;
            border: 1px solid {BORDER} !important;
            border-radius: 12px !important;
        }}

        .stSelectbox * {{
            color: {TEXT} !important;
        }}

        .stMultiSelect * {{
            color: {TEXT} !important;
        }}

        .stCheckbox label,
        .stRadio label,
        .stSelectbox label,
        .stTextInput label,
        .stTextArea label,
        .stDateInput label,
        .stTimeInput label,
        .stNumberInput label,
        .stFileUploader label,
        .stMultiSelect label,
        .stToggle label,
        .stMarkdown p,
        .stMarkdown li,
        .stMarkdown span,
        .stCaption,
        small {{
            color: {TEXT} !important;
            opacity: 1 !important;
        }}

        [data-testid="stCheckbox"] label,
        [data-testid="stRadio"] label,
        [data-testid="stWidgetLabel"] {{
            color: {TEXT} !important;
            opacity: 1 !important;
            font-weight: 500;
        }}

        [data-testid="stCheckbox"] div,
        [data-testid="stRadio"] div {{
            color: {TEXT} !important;
        }}

        .stAlert {{
            border-radius: 14px;
        }}

        .bloque-suave {{
            background-color: #ffffff;
            border: 1px solid {BORDER};
            border-radius: 14px;
            padding: 12px 16px;
            margin-bottom: 14px;
        }}
        </style>
        """,
        unsafe_allow_html=True
    )

aplicar_estilo()

# =============================
# LOGIN SIMPLE
# =============================
USER = "estudio"
PASSWORD = "peire2026"

if "logged_in" not in st.session_state:
    st.session_state.logged_in = False

if "menu_actual" not in st.session_state:
    st.session_state.menu_actual = "Dashboard"

def volver_al_dashboard():
    st.session_state.menu_actual = "Dashboard"

if not st.session_state.logged_in:
    st.markdown(
        """
        <div class="login-wrap">
            <div class="login-box">
                <div class="login-title">Estudio Peire</div>
                <div class="login-subtitle">Sistema interno de trabajo</div>
            </div>
        </div>
        """,
        unsafe_allow_html=True
    )

    username = st.text_input("Usuario", placeholder="Ingresá tu usuario")
    password = st.text_input("Contraseña", type="password", placeholder="Ingresá tu contraseña")

    if st.button("Ingresar al sistema"):
        if username == USER and password == PASSWORD:
            st.session_state.logged_in = True
            st.rerun()
        else:
            st.error("Usuario o contraseña incorrectos")

    st.stop()
# =============================
# HELPERS
# =============================
def exportar_word(texto: str, nombre_archivo: str):
    doc = Document()
    for linea in texto.strip().split("\n"):
        doc.add_paragraph(linea.rstrip())
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    st.download_button(
        label="⬇️ Descargar en Word",
        data=buffer,
        file_name=f"{nombre_archivo}.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )

def bloque_firma(firmante: str, matricula: str, estudio: str, contacto: str):
    out = "\n\n" + ("-" * 42) + "\n"
    out += f"Firma: {firmante.strip() if firmante else '_______________________'}\n"
    if matricula.strip():
        out += f"Matrícula: {matricula.strip()}\n"
    if estudio.strip():
        out += f"{estudio.strip()}\n"
    if contacto.strip():
        out += f"Contacto: {contacto.strip()}\n"
    return out

def safe(v: str, placeholder: str):
    v = (v or "").strip()
    return v if v else placeholder

def linea_amenaza(tono: str):
    if tono == "Neutral":
        return "bajo apercibimiento de iniciar las acciones legales que correspondan."
    if tono == "Firme":
        return "bajo apercibimiento de iniciar acciones legales sin más trámite, con más gastos y costas."
    return "bajo apercibimiento de promover de inmediato las acciones judiciales pertinentes, con más intereses, daños, gastos y costas."

def extraer_texto_archivo(uploaded_file):
    if uploaded_file is None:
        return ""

    nombre = uploaded_file.name.lower()

    try:
        if nombre.endswith(".txt"):
            uploaded_file.seek(0)
            return uploaded_file.read().decode("utf-8")

        elif nombre.endswith(".docx"):
            uploaded_file.seek(0)
            doc = Document(uploaded_file)
            texto = "\n".join([p.text for p in doc.paragraphs if p.text.strip()])
            return texto.strip()

        elif nombre.endswith(".pdf"):
            uploaded_file.seek(0)
            reader = PdfReader(uploaded_file)
            texto = ""

            for page in reader.pages:
                page_text = page.extract_text()
                if page_text:
                    texto += page_text + "\n"

            texto = texto.strip()

            if texto:
                return texto

            # Fallback OCR para PDF escaneado
            uploaded_file.seek(0)
            return extraer_texto_pdf_escaneado_con_ia(uploaded_file)

        elif nombre.endswith((".jpg", ".jpeg", ".png")):
            uploaded_file.seek(0)
            image_bytes = uploaded_file.read()

            if nombre.endswith(".jpg") or nombre.endswith(".jpeg"):
                mime_type = "image/jpeg"
            else:
                mime_type = "image/png"

            return extraer_texto_imagen_con_ia(image_bytes, mime_type=mime_type)

        else:
            return "ERROR_AL_LEER_ARCHIVO: Formato no soportado."

    except Exception as e:
        return f"ERROR_AL_LEER_ARCHIVO: {str(e)}"

def guardar_en_historial(tipo: str, titulo: str, contenido: str):
    if "historial_documentos" not in st.session_state:
        st.session_state["historial_documentos"] = []

    st.session_state["historial_documentos"].insert(0, {
        "tipo": tipo,
        "titulo": titulo,
        "contenido": contenido,
        "fecha": datetime.now().strftime("%d/%m/%Y %H:%M")
    })

def obtener_cliente_openai():
    api_key = None

    # Primero intenta desde Streamlit secrets (producción)
    if "OPENAI_API_KEY" in st.secrets:
        api_key = st.secrets["OPENAI_API_KEY"]
    
    # Si no, intenta desde variables de entorno (local)
    if not api_key:
        api_key = os.getenv("OPENAI_API_KEY")

    if not api_key:
        return None

    return OpenAI(api_key=api_key)

def generar_texto_con_ia(prompt_sistema: str, prompt_usuario: str):
    client = obtener_cliente_openai()

    if client is None:
        return "ERROR_IA: API KEY no configurada"

    try:
        respuesta = client.responses.create(
            model="gpt-4.1-mini",
            input=[
                {"role": "system", "content": prompt_sistema},
                {"role": "user", "content": prompt_usuario},
            ],
        )
        return (respuesta.output_text or "").strip()
    except Exception as e:
        return f"ERROR_IA: {str(e)}"

def diagnosticar_documento_con_ia(texto_documento: str, nombre_archivo: str = ""):
    client = obtener_cliente_openai()

    if client is None:
        return "ERROR_IA: API KEY no configurada"

    try:
        respuesta = client.responses.create(
            model="gpt-4.1-mini",
            input=[
                {
                    "role": "system",
                    "content": (
                        "Sos asistente jurídico del Estudio Peire. "
                        "Analizás documentos jurídicos en español argentino. "
                        "No inventes datos. Si algo no surge del documento, decí 'No detectado'. "
                        "Tu tarea es diagnosticar el documento y recomendar el próximo paso interno."
                    ),
                },
                {
                    "role": "user",
                    "content": f"""
Analizá el siguiente documento jurídico y devolvé el resultado con este formato exacto:

TIPO DETECTADO:
TEMA PRINCIPAL:
REMITENTE DETECTADO:
DESTINATARIO DETECTADO:
MONTO DETECTADO:
PLAZO DETECTADO:
RESUMEN EJECUTIVO:
RIESGO / URGENCIA:
PRÓXIMO PASO RECOMENDADO:
HERRAMIENTA DEL SISTEMA SUGERIDA:
ESTRATEGIA SUGERIDA:

Nombre del archivo: {nombre_archivo}

Texto del documento:
{texto_documento}
""",
                },
            ],
        )
        return (respuesta.output_text or "").strip()
    except Exception as e:
        return f"ERROR_IA: {str(e)}"

def editar_texto_con_ia(texto_original: str, instruccion_usuario: str):
    client = obtener_cliente_openai()

    if client is None:
        return "ERROR_IA: API KEY no configurada"

    try:
        respuesta = client.responses.create(
            model="gpt-4.1-mini",
            input=[
                {
                    "role": "system",
                    "content": (
                        "Sos asistente jurídico del Estudio Peire. "
                        "Tu tarea es editar y reescribir documentos jurídicos en español jurídico argentino. "
                        "No inventes hechos, normas ni jurisprudencia. "
                        "Respetá el contenido base y aplicá solo los cambios pedidos por el usuario. "
                        "Devolvé únicamente la nueva versión del texto completo."
                    ),
                },
                {
                    "role": "user",
                    "content": f"""
Texto original:
{texto_original}

Instrucción del usuario:
{instruccion_usuario}
""",
                },
            ],
        )
        return (respuesta.output_text or "").strip()
    except Exception as e:
        return f"ERROR_IA: {str(e)}"

def limpiar_resultado(clave: str):
    if clave in st.session_state:
        del st.session_state[clave]

def imagen_a_data_url(image_bytes, mime_type="image/png"):
    base64_image = base64.b64encode(image_bytes).decode("utf-8")
    return f"data:{mime_type};base64,{base64_image}"


def extraer_texto_imagen_con_ia(image_bytes, mime_type="image/png"):
    client = obtener_cliente_openai()

    if client is None:
        return "ERROR_OCR_IMAGEN: API KEY no configurada"

    try:
        data_url = imagen_a_data_url(image_bytes, mime_type)

        respuesta = client.responses.create(
            model="gpt-4.1-mini",
            input=[
                {
                    "role": "user",
                    "content": [
                        {
                            "type": "input_text",
                            "text": (
                                "Extraé todo el texto visible de esta imagen en español, "
                                "respetando lo más posible saltos de línea, nombres, fechas, montos y domicilios. "
                                "No resumas. No expliques. Devolvé solo el texto extraído."
                            ),
                        },
                        {
                            "type": "input_image",
                            "image_url": data_url,
                            "detail": "high",
                        },
                    ],
                }
            ],
        )

        return (respuesta.output_text or "").strip()

    except Exception as e:
        return f"ERROR_OCR_IMAGEN: {str(e)}"
    
def extraer_texto_pdf_escaneado_con_ia(uploaded_file):
    client = obtener_cliente_openai()

    if client is None:
        return "ERROR_OCR_PDF: API KEY no configurada"

    try:
        uploaded_file.seek(0)
        pdf_bytes = uploaded_file.read()

        doc = fitz.open(stream=pdf_bytes, filetype="pdf")
        textos_paginas = []

        for i, page in enumerate(doc):
            pix = page.get_pixmap(matrix=fitz.Matrix(2, 2), alpha=False)
            image_bytes = pix.tobytes("png")

            texto_pagina = extraer_texto_imagen_con_ia(image_bytes, mime_type="image/png")

            if texto_pagina.startswith("ERROR_OCR_IMAGEN:"):
                return texto_pagina

            if texto_pagina.strip():
                textos_paginas.append(f"\n--- Página {i + 1} ---\n{texto_pagina}")

        return "\n".join(textos_paginas).strip()

    except Exception as e:
        return f"ERROR_OCR_PDF: {str(e)}"

def imagen_a_data_url(image_bytes, mime_type="image/png"):
    base64_image = base64.b64encode(image_bytes).decode("utf-8")
    return f"data:{mime_type};base64,{base64_image}"

def extraer_datos_clave_con_ia(texto_documento: str):
    client = obtener_cliente_openai()

    if client is None:
        return {"error": "ERROR_IA: API KEY no configurada"}

    try:
        respuesta = client.responses.create(
            model="gpt-4.1-mini",
            input=[
                {
                    "role": "system",
                    "content": (
                        "Sos asistente jurídico del Estudio Peire. "
                        "Tu tarea es extraer datos clave de documentos jurídicos argentinos. "
                        "No inventes datos. Si un dato no surge claramente, devolvé 'No detectado'. "
                        "Respondé únicamente en JSON válido, sin markdown, sin bloque ```json, sin texto antes ni después."
                    ),
                },
                {
                    "role": "user",
                    "content": f"""
Extraé del siguiente documento estos campos exactos y devolvelos en JSON válido:

{{
  "tipo_documento": "",
  "remitente": "",
  "destinatario": "",
  "fecha": "",
  "monto": "",
  "objeto": "",
  "resumen": ""
}}

Importante:
- En tipo_documento elegí una de estas opciones exactas:
  "Carta Documento recibida"
  "Respuesta a Carta Documento"
  "Oficio recibido"
  "Intimación"
  "Otro"

Texto del documento:
{texto_documento}
""",
                },
            ],
        )

        texto_respuesta = (respuesta.output_text or "").strip()

        texto_respuesta = re.sub(r"^```json\s*", "", texto_respuesta, flags=re.IGNORECASE)
        texto_respuesta = re.sub(r"^```\s*", "", texto_respuesta)
        texto_respuesta = re.sub(r"\s*```$", "", texto_respuesta)

        return json.loads(texto_respuesta)

    except Exception as e:
        return {"error": f"ERROR_IA: {str(e)}"}

# =============================
# HEADER
# =============================
st.markdown(
    """
    <div class="bloque-header">
        <div class="titulo-principal">Estudio Peire</div>
        <div class="subtitulo-principal">Sistema interno de trabajo · uso exclusivo del estudio</div>
    </div>
    """,
    unsafe_allow_html=True
)
# =============================
# SIDEBAR
# =============================
opciones_menu = [
    "Dashboard",
    "Diagnóstico Inteligente",
    "Carta Documento",
    "Respuesta Carta Documento",
    "Contestación de Oficio",
    "Mailing (Modo Agente)",
    "Presupuesto",
    "Análisis de Documento",
    "Historial",
    "Biblioteca Oficial de Prompts",
]

menu = st.sidebar.radio(
    "Herramientas",
    opciones_menu,
    index=opciones_menu.index(st.session_state.menu_actual),
)

st.session_state.menu_actual = menu

st.sidebar.markdown("---")

with st.sidebar.expander("🖊️ Datos de firma (se agregan al final)", expanded=False):
    firmante = st.text_input("Firmante", value="")
    matricula = st.text_input("Matrícula (opcional)", value="")
    estudio = st.text_input("Nombre del estudio", value="Estudio Peire")
    contacto = st.text_input("Contacto (email/teléfono)", value="")

if st.sidebar.button("Cerrar sesión"):
    st.session_state.logged_in = False
    st.rerun()
# =========================================================
# DASHBOARD
# =========================================================
if menu == "Dashboard":
    st.markdown("## Panel principal")
    st.markdown("Elegí la tarea que querés realizar.")

    col1, col2 = st.columns(2)

    with col1:
        st.markdown(
            """
            <div class="mini-card">
                <h4>📄 Crear Carta Documento</h4>
                <p>Generar intimaciones, reclamos y borradores formales.</p>
            </div>
            """,
            unsafe_allow_html=True
        )
        if st.button("Ir a Carta Documento", key="go_cd"):
            st.session_state.menu_actual = "Carta Documento"
            st.rerun()

        st.markdown(
            """
            <div class="mini-card">
                <h4>📑 Contestar Oficio</h4>
                <p>Redactar contestaciones formales de oficio.</p>
            </div>
            """,
            unsafe_allow_html=True
        )
        if st.button("Ir a Contestación de Oficio", key="go_oficio"):
            st.session_state.menu_actual = "Contestación de Oficio"
            st.rerun()

        st.markdown(
            """
            <div class="mini-card">
                <h4>📧 Redactar Mailing</h4>
                <p>Preparar correos y comunicaciones con clientes.</p>
            </div>
            """,
            unsafe_allow_html=True
        )
        if st.button("Ir a Mailing", key="go_mail"):
            st.session_state.menu_actual = "Mailing (Modo Agente)"
            st.rerun()

        st.markdown(
            """
            <div class="mini-card">
                <h4>📂 Analizar Documento</h4>
                <p>Subir archivos, extraer texto y preparar un análisis base.</p>
            </div>
            """,
            unsafe_allow_html=True
        )
        if st.button("Ir a Análisis de Documento", key="go_analisis"):
            st.session_state.menu_actual = "Análisis de Documento"
            st.rerun()

    with col2:
        st.markdown(
            """
            <div class="mini-card">
                <h4>✉️ Responder Carta Documento</h4>
                <p>Preparar respuestas a documentos recibidos.</p>
            </div>
            """,
            unsafe_allow_html=True
        )
        if st.button("Ir a Respuesta Carta Documento", key="go_resp"):
            st.session_state.menu_actual = "Respuesta Carta Documento"
            st.rerun()

        st.markdown(
            """
            <div class="mini-card">
                <h4>🧠 Diagnóstico Inteligente</h4>
                <p>Clasificar documentos, detectar datos clave y sugerir el próximo paso.</p>
            </div>
            """,
            unsafe_allow_html=True
        )
        if st.button("Ir a Diagnóstico Inteligente", key="go_diag"):
            st.session_state.menu_actual = "Diagnóstico Inteligente"
            st.rerun()

        st.markdown(
            """
            <div class="mini-card">
                <h4>💼 Hacer Presupuesto</h4>
                <p>Generar propuestas y presupuestos de honorarios.</p>
            </div>
            """,
            unsafe_allow_html=True
        )
        if st.button("Ir a Presupuesto", key="go_presupuesto"):
            st.session_state.menu_actual = "Presupuesto"
            st.rerun()

        st.markdown(
            """
            <div class="mini-card">
                <h4>📚 Ver Prompts</h4>
                <p>Acceder a la biblioteca interna de prompts del estudio.</p>
            </div>
            """,
            unsafe_allow_html=True
        )
        if st.button("Ir a Biblioteca de Prompts", key="go_prompts"):
            st.session_state.menu_actual = "Biblioteca Oficial de Prompts"
            st.rerun()

# =========================================================
# DIAGNÓSTICO INTELIGENTE
# =========================================================
elif menu == "Diagnóstico Inteligente":
    st.header("🧠 Diagnóstico Inteligente de Documento")

    col_a, col_b = st.columns([1, 1])
    with col_a:
        st.button("← Volver al panel principal", on_click=volver_al_dashboard)
    with col_b:
        if st.button("Nuevo diagnóstico", key="reset_diagnostico"):
            limpiar_resultado("ultimo_diagnostico")
            limpiar_resultado("editor_diagnostico")
            limpiar_resultado("sync_editor_diagnostico")
            limpiar_resultado("instruccion_edicion_diagnostico")
            st.rerun()
    
    st.write("Subí un documento y la IA va a detectar automáticamente qué es, de qué trata y cuál sería el próximo paso recomendado dentro del sistema.")

    archivo_diag = st.file_uploader(
    "Subir documento para diagnóstico",
    type=["pdf", "docx", "txt", "jpg", "jpeg", "png"],
    key="archivo_diagnostico"
    )

    observaciones_diag = st.text_area(
        "Observaciones del estudio (opcional)",
        height=100,
        placeholder="Ej: llegó hoy, parece reclamo por alquiler, cliente dice que no corresponde, etc."
    )

    texto_diagnostico = ""

    if archivo_diag is not None:
        with st.spinner("Procesando archivo..."):
            texto_diagnostico = extraer_texto_archivo(archivo_diag)

        if texto_diagnostico.startswith("ERROR_"):
            st.error(texto_diagnostico)
            texto_diagnostico = ""

        elif texto_diagnostico.strip():
            st.success(f"Archivo cargado: {archivo_diag.name}")

            st.text_area(
                "Texto detectado del archivo",
                value=texto_diagnostico,
                height=220,
                key="texto_detectado_diagnostico"
            )

        else:
            st.warning("No se pudo extraer texto del archivo o está vacío.")

        # 👇 ESTO TAMBIÉN ADENTRO
        if texto_diagnostico == "PDF_ESCANEADO_O_SIN_TEXTO":
            st.warning("El PDF no tiene texto extraíble. Parece ser un documento escaneado.")

        elif texto_diagnostico == "IMAGEN_CARGADA_PARA_OCR":
            st.warning("Se cargó una imagen. Falta agregar OCR.")

        elif texto_diagnostico.startswith("ERROR_AL_LEER_ARCHIVO:"):
            st.error(texto_diagnostico)

        elif texto_diagnostico.strip():
            st.success("Archivo leído correctamente.")
    
    if st.button("Generar diagnóstico con IA"):
        
        limpiar_resultado("ultimo_diagnostico")
        if not texto_diagnostico.strip():
            st.warning("Primero subí un archivo válido para diagnosticar.")
        else:
            texto_base = texto_diagnostico
            if observaciones_diag.strip():
                texto_base += f"\n\nObservaciones del estudio:\n{observaciones_diag}"

            diagnostico = diagnosticar_documento_con_ia(
                texto_documento=texto_base,
                nombre_archivo=archivo_diag.name if archivo_diag else ""
            )

            if not diagnostico:
                st.error("No se encontró OPENAI_API_KEY en Secrets.")
                st.stop()

            if str(diagnostico).startswith("ERROR_IA:"):
                st.error(diagnostico)
                st.stop()

            st.session_state["ultimo_diagnostico"] = diagnostico
            st.session_state["sync_editor_diagnostico"] = True

            guardar_en_historial(
                tipo="Diagnóstico Inteligente",
                titulo=f"Diagnóstico - {archivo_diag.name if archivo_diag else 'Sin archivo'}",
                contenido=diagnostico
            )

    if "ultimo_diagnostico" in st.session_state:

        if "editor_diagnostico" not in st.session_state:
            st.session_state["editor_diagnostico"] = st.session_state["ultimo_diagnostico"]

        if st.session_state.get("sync_editor_diagnostico", False):
            st.session_state["editor_diagnostico"] = st.session_state["ultimo_diagnostico"]
            st.session_state["sync_editor_diagnostico"] = False

        st.markdown("### Resultado del diagnóstico")

        texto_actual_diagnostico = st.text_area(
            "Diagnóstico generado / editable",
            height=420,
            key="editor_diagnostico"
        )

        st.session_state["ultimo_diagnostico"] = texto_actual_diagnostico

        st.markdown("### Editar diagnóstico con IA")
        instruccion_diag = st.text_input(
            "Pedile cambios a la IA",
            value=st.session_state.get("instruccion_edicion_diagnostico", ""),
            placeholder="Ej: resumilo más, agregá más detalle en la estrategia, hacelo más claro."
        )

        if st.button("Aplicar cambios al diagnóstico con IA", key="editar_diagnostico_ia"):
            if not instruccion_diag.strip():
                st.warning("Escribí una instrucción para editar el diagnóstico.")
            else:
                texto_editado_diag = editar_texto_con_ia(
                    texto_actual_diagnostico,
                    instruccion_diag
                )

                if not texto_editado_diag:
                    st.error("No se encontró OPENAI_API_KEY en Secrets.")
                elif str(texto_editado_diag).startswith("ERROR_IA:"):
                    st.error(texto_editado_diag)
                else:
                    st.session_state["ultimo_diagnostico"] = texto_editado_diag
                    st.session_state["sync_editor_diagnostico"] = True

                    guardar_en_historial(
                        tipo="Edición IA - Diagnóstico Inteligente",
                        titulo="Edición IA - Diagnóstico",
                        contenido=texto_editado_diag
                    )

                    st.success("Diagnóstico actualizado con IA.")
                    st.rerun()

        exportar_word(
            st.session_state["ultimo_diagnostico"],
            "Diagnostico_Inteligente_Estudio_Peire"
        )
# =========================================================
# 1) CARTA DOCUMENTO
# =========================================================
elif menu == "Carta Documento":
    st.header("📄 Carta Documento")

    col_a, col_b = st.columns([1, 1])
    with col_a:
        st.button("← Volver al panel principal", on_click=volver_al_dashboard, key="volver_carta_doc")
    with col_b:
        if st.button("Nuevo documento", key="reset_carta"):
            limpiar_resultado("ultimo_texto_carta_documento")
            limpiar_resultado("editor_carta_documento")
            limpiar_resultado("sync_editor_carta_documento")
            limpiar_resultado("instruccion_edicion_carta_documento")
            st.rerun()

    col1, col2 = st.columns(2)
    with col1:
        remitente = st.text_input("Remitente / Cliente", placeholder="Ej: Carlos Raúl Fernández")
        dom_remitente = st.text_input("Domicilio remitente", placeholder="Ej: Av. Santa Fe 2450, CABA")
    with col2:
        destinatario = st.text_input("Destinatario", placeholder="Ej: Juan Pérez")
        dom_destinatario = st.text_input("Domicilio destinatario", placeholder="Ej: Av. Rivadavia 1234, CABA")

    col3, col4, col5 = st.columns(3)
    with col3:
        jurisd = st.text_input("Jurisdicción / Ciudad", value="CABA", placeholder="Ej: CABA")
    with col4:
        fecha = st.text_input("Fecha (dd/mm/aaaa)", value=date.today().strftime("%d/%m/%Y"), placeholder="Ej: 06/03/2026")
    with col5:
        plazo = st.selectbox("Plazo que se intima", ["24 hs", "48 hs", "72 hs", "5 días", "10 días", "15 días"])

    tipo = st.selectbox(
        "Tipo",
        [
            "Intimación de pago (deuda)",
            "Intimación por incumplimiento (cumplimiento de obligación)",
            "Rescisión / Resolución contractual",
            "Cese de conducta / daños",
            "Laboral (intimación / regularización)",
            "Otra (personalizada)",
        ],
    )

    col6, col7 = st.columns(2)
    with col6:
        monto = st.text_input("Monto (si aplica)", placeholder="Ej: $450.000")
    with col7:
        referencia = st.text_input("Referencia/Contrato/Expte (opcional)", placeholder="Ej: Contrato 01/12/2025")

    hechos = st.text_area(
        "Hechos / Antecedentes (cronología breve)",
        height=120,
        placeholder="Describí brevemente los hechos en orden cronológico."
    )

    pedido_concreto = st.text_area(
        "Pedido concreto (qué exigís que haga la otra parte)",
        height=90,
        placeholder="Ej: Intimo a abonar la suma adeudada dentro del plazo indicado."
    )

    col8, col9, col10, col11 = st.columns(4)
    with col8:
        mencionar_pruebas = st.checkbox("Mencionar documentación/pruebas", value=True)
    with col9:
        incluir_reserva = st.checkbox("Reserva de acciones y derechos", value=True)
    with col10:
        incluir_costas = st.checkbox("Apercibimiento de gastos y costas", value=True)
    with col11:
        abrir_acuerdo = st.checkbox("Abrir posibilidad de acuerdo", value=False)

    usar_ia = st.checkbox("Usar IA para redactar la carta", value=True)

    texto_personalizado = ""
    if tipo == "Otra (personalizada)":
        texto_personalizado = st.text_area(
            "Texto base personalizado (1–4 líneas)",
            height=80,
            placeholder="Escribí una base personalizada para esta carta documento."
        )

    if st.button("Generar Carta Documento", key="generar_carta_doc"):
        limpiar_resultado("ultimo_texto_carta_documento")

        if usar_ia:
            prompt_sistema = (
                "Sos asistente jurídico del Estudio Peire. "
                "Redactás CARTAS DOCUMENTO en español jurídico argentino, con tono profesional y prudente. "
                "No inventes hechos, normas ni jurisprudencia. "
                "No cites artículos si no fueron dados por el usuario. "
                "Devolvé directamente el texto final del documento listo para revisar por un abogado."
            )

            prompt_usuario = f"""
Redactá una CARTA DOCUMENTO en español jurídico argentino.

Tipo: {tipo}
Tono: formal
Lugar/Jurisdicción: {jurisd}
Fecha: {fecha}
Plazo intimado: {plazo}

Remitente: {remitente}
Domicilio remitente: {dom_remitente}

Destinatario: {destinatario}
Domicilio destinatario: {dom_destinatario}

Monto: {monto}
Referencia/Contrato/Expte: {referencia}

Hechos / antecedentes:
{hechos}

Pedido concreto:
{pedido_concreto}

Mencionar documentación/pruebas: {mencionar_pruebas}
Reserva de acciones y derechos: {incluir_reserva}
Apercibimiento de gastos y costas: {incluir_costas}
Abrir posibilidad de acuerdo: {abrir_acuerdo}

Texto base personalizado:
{texto_personalizado}

Firmante: {firmante}
Matrícula: {matricula}
Estudio: {estudio}
Contacto: {contacto}

Devolvé solo el texto final del documento, sin explicaciones adicionales.
"""
            t = generar_texto_con_ia(prompt_sistema, prompt_usuario)

            if not t:
                st.error("No se encontró OPENAI_API_KEY en Secrets.")
                st.stop()

            if str(t).startswith("ERROR_IA:"):
                st.error(t)
                st.stop()

            if incluir_reserva and "reserva" not in t.lower():
                t += "\n\nSe reserva expresamente el ejercicio de acciones y derechos."

            if incluir_costas and "gastos y costas" not in t.lower():
                t += "\n\nTodo ello con más intereses, gastos y costas."

            t += bloque_firma(firmante, matricula, estudio, contacto)

        else:
            t = "CARTA DOCUMENTO\n"
            t += f"Lugar/Jurisdicción: {safe(jurisd, '[Lugar]')}\n"
            t += f"Fecha: {safe(fecha, '[Fecha]')}\n"
            if referencia.strip():
                t += f"Referencia: {referencia.strip()}\n"
            t += "\n"
            t += f"Remitente: {safe(remitente, '[Remitente]')}\n"
            t += f"Domicilio: {safe(dom_remitente, '[Domicilio remitente]')}\n"
            t += f"Destinatario: {safe(destinatario, '[Destinatario]')}\n"
            t += f"Domicilio: {safe(dom_destinatario, '[Domicilio destinatario]')}\n"

            t += "\n\nPor la presente, "

            if tipo == "Intimación de pago (deuda)":
                t += f"INTIMO a Ud. para que en el plazo de {plazo} abone la suma de {safe(monto, '[Monto]')} en concepto de deuda, {linea_amenaza('Firme')}"
            elif tipo == "Intimación por incumplimiento (cumplimiento de obligación)":
                t += f"INTIMO a Ud. para que en el plazo de {plazo} cumpla íntegramente con lo debido, {linea_amenaza('Firme')}"
            elif tipo == "Rescisión / Resolución contractual":
                t += f"INTIMO a Ud. para que en el plazo de {plazo} regularice su situación contractual, bajo apercibimiento de considerar resuelto el vínculo y reclamar daños, {linea_amenaza('Firme')}"
            elif tipo == "Cese de conducta / daños":
                t += f"INTIMO a Ud. para que en el plazo de {plazo} cese la conducta lesiva denunciada y adopte las medidas necesarias, {linea_amenaza('Firme')}"
            elif tipo == "Laboral (intimación / regularización)":
                t += f"INTIMO a Ud. para que en el plazo de {plazo} regularice la situación denunciada, {linea_amenaza('Firme')}"
            else:
                base = safe(texto_personalizado, "INTIMO a Ud. para que en el plazo indicado cumpla con lo requerido,")
                t += f"{base} {linea_amenaza('Firme')}"

            t += "\n\nHechos/antecedentes:\n"
            t += safe(hechos, "[Describir hechos en forma breve y cronológica]")

            if pedido_concreto.strip():
                t += "\n\nPedido concreto:\n"
                t += pedido_concreto.strip()

            if mencionar_pruebas:
                t += "\n\nSe deja constancia que existen antecedentes y/o documentación respaldatoria que acreditan lo aquí expuesto."
            if abrir_acuerdo:
                t += "\n\nSin perjuicio de lo anterior, se deja abierta la posibilidad de arribar a una solución consensuada en términos razonables."
            if incluir_costas:
                t += "\n\nTodo ello con más intereses, gastos y costas."
            if incluir_reserva:
                t += "\n\nSe reserva expresamente el ejercicio de acciones y derechos."

            t += "\n\nQueda Ud. debidamente notificado."
            t += bloque_firma(firmante, matricula, estudio, contacto)

        st.session_state["ultimo_texto_carta_documento"] = t
        st.session_state["sync_editor_carta_documento"] = True

        guardar_en_historial(
            tipo="Carta Documento",
            titulo=f"Carta Documento - {destinatario or 'Sin destinatario'}",
            contenido=t
        )

        st.rerun()

    if "ultimo_texto_carta_documento" in st.session_state:

        if "editor_carta_documento" not in st.session_state:
            st.session_state["editor_carta_documento"] = st.session_state["ultimo_texto_carta_documento"]

        if st.session_state.get("sync_editor_carta_documento", False):
            st.session_state["editor_carta_documento"] = st.session_state["ultimo_texto_carta_documento"]
            st.session_state["sync_editor_carta_documento"] = False

        st.markdown("### Resultado")

        texto_actual_cd = st.text_area(
            "Texto generado / editable",
            height=420,
            key="editor_carta_documento"
        )

        st.session_state["ultimo_texto_carta_documento"] = texto_actual_cd

        st.markdown("### Editar con IA")
        instruccion_edicion_cd = st.text_input(
            "Pedile cambios a la IA",
            value=st.session_state.get("instruccion_edicion_carta_documento", ""),
            placeholder="Ej: hacelo más firme, más técnico, más breve."
        )

        if st.button("Aplicar cambios con IA", key="editar_carta_ia"):
            if not instruccion_edicion_cd.strip():
                st.warning("Escribí una instrucción para editar el texto.")
            else:
                texto_editado_cd = editar_texto_con_ia(texto_actual_cd, instruccion_edicion_cd)

                if not texto_editado_cd:
                    st.error("No se encontró OPENAI_API_KEY en Secrets.")
                elif str(texto_editado_cd).startswith("ERROR_IA:"):
                    st.error(texto_editado_cd)
                else:
                    st.session_state["ultimo_texto_carta_documento"] = texto_editado_cd
                    st.session_state["sync_editor_carta_documento"] = True

                    guardar_en_historial(
                        tipo="Edición IA - Carta Documento",
                        titulo=f"Edición IA - {destinatario or 'Sin destinatario'}",
                        contenido=texto_editado_cd
                    )

                    st.success("Texto actualizado con IA.")
                    st.rerun()

        exportar_word(
            st.session_state["ultimo_texto_carta_documento"],
            "Carta_Documento_Estudio_Peire"
        )

# =========================================================
# 2) RESPUESTA A CARTA DOCUMENTO
# =========================================================
elif menu == "Respuesta Carta Documento":
    st.header("✉️ Respuesta a Carta Documento")

    col_a, col_b = st.columns([1, 1])
    with col_a:
        st.button("← Volver al panel principal", on_click=volver_al_dashboard)
    with col_b:
        if st.button("Nueva respuesta", key="reset_respuesta_cd"):
            limpiar_resultado("ultimo_texto_respuesta_cd")
            limpiar_resultado("editor_respuesta_cd")
            limpiar_resultado("sync_editor_respuesta_cd")
            limpiar_resultado("instruccion_edicion_respuesta_cd")
            st.rerun()
    
    datos_analisis = st.session_state.get("analisis_para_respuesta", {})

    if datos_analisis:

        st.success("Se cargó información desde 'Análisis de Documento'.")

        col_ref1, col_ref2 = st.columns(2)

        with col_ref1:
            remitente = st.text_input(
                "Remitente detectado",
                value=datos_analisis.get("remitente", ""),
                key="remitente_detectado"
            )

            destinatario = st.text_input(
                "Destinatario detectado",
                value=datos_analisis.get("destinatario", ""),
                key="destinatario_detectado"
            )

            tipo_documento = st.text_input(
                "Tipo de documento",
                value=datos_analisis.get("tipo_documento", ""),
                key="tipo_documento_detectado"
            )

        with col_ref2:
            fecha_doc = st.text_input(
                "Fecha detectada",
                value=datos_analisis.get("fecha_doc", ""),
                key="fecha_detectada"
            )

            monto = st.text_input(
                "Monto detectado",
                value=datos_analisis.get("monto", ""),
                key="monto_detectado"
            )

            objeto = st.text_input(
                "Objeto detectado",
                value=datos_analisis.get("objeto", ""),
                key="objeto_detectado"
            )

        if st.button("🧹 Limpiar datos cargados del análisis"):

            if "analisis_para_respuesta" in st.session_state:
                del st.session_state["analisis_para_respuesta"]

            st.rerun()

    col1, col2, col3 = st.columns(3)
    with col1:
        postura = st.selectbox("Postura", ["Negar deuda/hechos", "Aceptar parcialmente", "Proponer acuerdo", "Rechazar e intimar"])
    with col2:
        tono = st.selectbox("Tono", ["Neutral", "Firme", "Muy firme"])
    with col3:
        plazo_intimacion = st.selectbox("Si intimás, plazo", ["24 hs", "48 hs", "72 hs", "5 días", "10 días"])

    archivo_respuesta = st.file_uploader(
    "Subir documento recibido (opcional)",
    type=["pdf", "docx", "txt", "jpg", "jpeg", "png"],
    key="archivo_respuesta_cd"
    )   

    texto_archivo_respuesta = ""
    
    if archivo_respuesta is not None:
        with st.spinner("Procesando archivo..."):
            texto_archivo_respuesta = extraer_texto_archivo(archivo_respuesta)

        if texto_archivo_respuesta.startswith("ERROR_"):
            st.error(texto_archivo_respuesta)
            texto_archivo_respuesta = ""
        elif texto_archivo_respuesta.strip():
            st.success(f"Archivo cargado: {archivo_respuesta.name}")
            st.text_area(
                "Texto detectado del archivo",
                value=texto_archivo_respuesta,
                height=180,
                key="texto_detectado_respuesta"
            )
        else:
            st.warning("No se pudo extraer texto del archivo o está vacío.")

        if texto_archivo_respuesta == "PDF_ESCANEADO_O_SIN_TEXTO":
            st.warning("El PDF no tiene texto extraíble. Parece ser un documento escaneado o una imagen en PDF.")

        elif texto_archivo_respuesta == "IMAGEN_CARGADA_PARA_OCR":
            st.warning("Se cargó una imagen. Falta agregar lectura OCR para extraer el texto automáticamente.")

        elif texto_archivo_respuesta.startswith("ERROR_AL_LEER_ARCHIVO:"):
            st.error(texto_archivo_respuesta)

        else:
            st.success("Archivo leído correctamente.")
    
    texto_base_respuesta = ""
    if texto_archivo_respuesta.strip():
        texto_base_respuesta = texto_archivo_respuesta
    else:
        texto_base_respuesta = datos_analisis.get("texto_recibido", "")

    texto_recibido = st.text_area(
        "Texto recibido (pegar o editar)",
        value=texto_base_respuesta,
        height=160,
        placeholder="Pegá acá el texto recibido o cargá un archivo arriba."
    )

    hechos_reales = st.text_area(
        "Hechos reales del cliente (lo que SÍ pasó)",
        value=datos_analisis.get("hechos_reales", ""),
        height=120,
        placeholder="Describí la versión real del cliente."
    )

    col4, col5, col6, col7 = st.columns(4)
    with col4:
        mencionar_pruebas = st.checkbox("Mencionar pruebas/documentación", value=True)
    with col5:
        incluir_reserva = st.checkbox("Reserva de acciones y derechos", value=True)
    with col6:
        incluir_costas = st.checkbox("Apercibimiento de gastos y costas", value=True)
    with col7:
        intimar_cese = st.checkbox("Intimar rectificación / cese de reclamo", value=False)

    usar_ia = st.checkbox("Usar IA para mejorar la redacción", value=True)

    propuesta = ""
    if postura in ["Aceptar parcialmente", "Proponer acuerdo"]:
        propuesta = st.text_area(
            "Propuesta (pago/plan/condiciones)",
            height=90,
            placeholder="Ej: Se propone plan de pago en 3 cuotas mensuales."
        )

    if st.button("Generar Respuesta"):
        
        limpiar_resultado("ultimo_texto_respuesta_cd")    
        if usar_ia:
            prompt_sistema = (
                "Sos asistente jurídico del Estudio Peire. "
                "Redactás borradores claros, profesionales y prudentes en español jurídico argentino. "
                "No inventes normas, jurisprudencia ni hechos. "
                "No cites artículos si no fueron dados por el usuario. "
                "Redactá un texto listo para revisar por un abogado."
            )

            prompt_usuario = f"""
Redactá una RESPUESTA A CARTA DOCUMENTO en español jurídico argentino.

Postura: {postura}
Tono: {tono}
Plazo de intimación si corresponde: {plazo_intimacion}

Texto recibido:
{texto_recibido}

Hechos reales del cliente:
{hechos_reales}

Mencionar pruebas/documentación: {mencionar_pruebas}
Reserva de acciones y derechos: {incluir_reserva}
Apercibimiento de gastos y costas: {incluir_costas}
Intimar rectificación / cese de reclamo: {intimar_cese}

Propuesta:
{propuesta if propuesta else "No hay propuesta."}

Firmante: {firmante}
Matrícula: {matricula}
Estudio: {estudio}
Contacto: {contacto}

Devolvé solo el texto final del documento, sin explicaciones adicionales.
"""
            t = generar_texto_con_ia(prompt_sistema, prompt_usuario)

            if not t:
                st.error("No se encontró OPENAI_API_KEY en Secrets.")
                st.stop()

            if str(t).startswith("ERROR_IA:"):
                st.error(t)
                st.stop()

            if incluir_reserva and "reserva" not in t.lower():
                t += "\n\nSe reserva expresamente el ejercicio de acciones y derechos."

            if incluir_costas and "gastos y costas" not in t.lower():
                t += "\n\nTodo ello con más gastos y costas."

            t += bloque_firma(firmante, matricula, estudio, contacto)

        else:
            t = "RESPUESTA A CARTA DOCUMENTO\n\n"
            t += "En relación a su comunicación, mediante la cual manifiesta:\n\n"
            t += safe(texto_recibido, "[Pegar texto recibido]") + "\n\n"

            if postura == "Negar deuda/hechos":
                if tono == "Neutral":
                    t += "Se rechazan los hechos y manifestaciones allí vertidas por no ajustarse a la realidad.\n"
                elif tono == "Firme":
                    t += "Se rechazan los hechos y el derecho invocados por improcedentes y carentes de sustento.\n"
                else:
                    t += "Se niegan categóricamente los hechos y el derecho invocados por resultar falsos, improcedentes y carentes de respaldo.\n"
            elif postura == "Aceptar parcialmente":
                t += "Se efectúan las siguientes aclaraciones, aceptándose únicamente lo que se indica en forma expresa y rechazándose todo lo demás.\n"
            elif postura == "Proponer acuerdo":
                t += "Sin reconocer hechos ni derecho y a fin de evitar mayores costos y litigiosidad, se propone la siguiente vía de solución.\n"
            else:
                if tono == "Muy firme":
                    t += "Se rechazan de plano sus manifestaciones y se lo intima a cesar con reclamos infundados.\n"
                else:
                    t += "Se rechazan sus manifestaciones y se lo intima a adecuar su conducta conforme derecho.\n"

            t += "\nHechos reales / posición de mi representado:\n"
            t += safe(hechos_reales, "[Describir hechos reales]") + "\n"

            if mencionar_pruebas:
                t += "\nSe deja constancia que se cuenta con documentación y/o elementos probatorios respaldatorios, los cuales serán oportunamente acompañados de corresponder.\n"

            if propuesta.strip():
                t += "\nPropuesta:\n" + propuesta.strip() + "\n"

            if intimar_cese:
                t += f"\nINTIMO a Ud. a rectificar y/o cesar el reclamo improcedente en el plazo de {plazo_intimacion}, bajo apercibimiento de iniciar las acciones pertinentes.\n"

            if incluir_costas:
                t += "\nTodo ello con más gastos y costas.\n"
            if incluir_reserva:
                t += "\nSe reserva expresamente el ejercicio de acciones y derechos.\n"

            t += "\nQueda Ud. debidamente notificado.\n"
            t += bloque_firma(firmante, matricula, estudio, contacto)

        st.session_state["ultimo_texto_respuesta_cd"] = t
        st.session_state["sync_editor_respuesta_cd"] = True

        guardar_en_historial(
            tipo="Respuesta Carta Documento",
            titulo=f"Respuesta CD - {datos_analisis.get('remitente', 'Sin remitente')}",
            contenido=t
        )

    if "ultimo_texto_respuesta_cd" in st.session_state:

        if "editor_respuesta_cd" not in st.session_state:
            st.session_state["editor_respuesta_cd"] = st.session_state["ultimo_texto_respuesta_cd"]

        if st.session_state.get("sync_editor_respuesta_cd", False):
            st.session_state["editor_respuesta_cd"] = st.session_state["ultimo_texto_respuesta_cd"]
            st.session_state["sync_editor_respuesta_cd"] = False

        st.markdown("### Resultado")

        texto_actual = st.text_area(
            "Texto generado / editable",
            height=420,
            key="editor_respuesta_cd"
        )

        st.session_state["ultimo_texto_respuesta_cd"] = texto_actual

        st.markdown("### Editar con IA")
        instruccion_edicion = st.text_input(
            "Pedile cambios a la IA",
            value=st.session_state.get("instruccion_edicion_respuesta_cd", ""),
            placeholder="Ej: hacelo más firme, más corto, agregá reserva de daños."
        )

        if st.button("Aplicar cambios con IA", key="editar_respuesta_cd_ia"):
            if not instruccion_edicion.strip():
                st.warning("Escribí una instrucción para editar el texto.")
            else:
                texto_editado = editar_texto_con_ia(texto_actual, instruccion_edicion)

                if not texto_editado:
                    st.error("No se encontró OPENAI_API_KEY en Secrets.")
                elif str(texto_editado).startswith("ERROR_IA:"):
                    st.error(texto_editado)
                else:
                    st.session_state["ultimo_texto_respuesta_cd"] = texto_editado
                    st.session_state["sync_editor_respuesta_cd"] = True

                    guardar_en_historial(
                        tipo="Edición IA - Respuesta Carta Documento",
                        titulo=f"Edición IA - {datos_analisis.get('remitente', 'Sin remitente')}",
                        contenido=texto_editado
                    )

                    st.success("Texto actualizado con IA.")
                    st.rerun()

        exportar_word(
            st.session_state["ultimo_texto_respuesta_cd"],
            "Respuesta_Carta_Documento_Estudio_Peire"
        )

# =========================================================
# 3) CONTESTACIÓN DE OFICIO
# =========================================================
elif menu == "Contestación de Oficio":
    st.header("📑 Contestación de Oficio")

    col_a, col_b = st.columns([1, 1])
    with col_a:
        st.button("← Volver al panel principal", on_click=volver_al_dashboard)
    with col_b:
        if st.button("Nueva contestación", key="reset_oficio"):
            limpiar_resultado("ultimo_oficio")
            limpiar_resultado("editor_oficio")
            limpiar_resultado("sync_editor_oficio")
            limpiar_resultado("instruccion_edicion_oficio")
            st.rerun()
    
    col1, col2 = st.columns(2)
    with col1:
        organismo = st.text_input("Organismo / Juzgado")
        dependencia = st.text_input("Dependencia/Secretaría (opcional)")
    with col2:
        expediente = st.text_input("Carátula / Expediente")
        fecha = st.text_input("Fecha (dd/mm/aaaa)", value=date.today().strftime("%d/%m/%Y"))

    objeto = st.text_input("Objeto del oficio (1 línea)")
    pedido = st.text_area("Pedido del oficio (copiar/pegar)", height=110)
    respuesta = st.text_area("Información a informar (ordenada y completa)", height=140)

    col3, col4, col5 = st.columns(3)
    with col3:
        confidencialidad = st.checkbox("Agregar nota de confidencialidad/uso exclusivo", value=True)
    with col4:
        adjuntos = st.text_input("Adjuntos (listar, opcional)", value="")
    with col5:
        requiere_consent = st.checkbox("Aclarar facultades/consentimiento (si aplica)", value=False)

    usar_ia_oficio = st.checkbox("Usar IA para redactar la contestación", value=True)
    
    if st.button("Generar Contestación"):

        limpiar_resultado("ultimo_oficio")
        
        if usar_ia_oficio:
            prompt_sistema = (
                "Sos asistente jurídico del Estudio Peire. "
                "Redactás contestaciones de oficio en español jurídico argentino. "
                "El texto debe ser formal, claro, ordenado y apto para revisión profesional. "
                "No inventes hechos, normas ni jurisprudencia. "
                "Devolvé solo el texto final del documento."
            )

            prompt_usuario = f"""
Redactá una CONTESTACIÓN DE OFICIO en español jurídico argentino.

Organismo / Juzgado: {organismo}
Dependencia / Secretaría: {dependencia}
Expediente: {expediente}
Fecha: {fecha}
Objeto: {objeto}

Pedido del oficio:
{pedido}

Información a informar:
{respuesta}

Adjuntos: {adjuntos}
Incluir confidencialidad: {confidencialidad}
Requiere aclaración de consentimiento/facultades: {requiere_consent}

Firmante: {firmante}
Matrícula: {matricula}
Estudio: {estudio}
Contacto: {contacto}

Devolvé solo el texto final del documento.
"""
            t = generar_texto_con_ia(prompt_sistema, prompt_usuario)

            if not t:
                st.error("No se encontró OPENAI_API_KEY en Secrets.")
                st.stop()

            if str(t).startswith("ERROR_IA:"):
                st.error(t)
                st.stop()

            t += bloque_firma(firmante, matricula, estudio, contacto)

        else:
            t = "CONTESTACIÓN DE OFICIO\n\n"
            t += f"A: {safe(organismo,'[Organismo/Juzgado]')}\n"
            if dependencia.strip():
                t += f"Dependencia: {dependencia.strip()}\n"
            if expediente.strip():
                t += f"Ref.: {expediente.strip()}\n"
            t += f"Fecha: {safe(fecha,'[Fecha]')}\n\n"

            if objeto.strip():
                t += f"Objeto: {objeto.strip()}\n\n"

            t += "En respuesta al oficio recibido, se informa lo siguiente:\n\n"
            if pedido.strip():
                t += "I. Pedido del oficio:\n"
                t += pedido.strip() + "\n\n"

            t += "II. Respuesta:\n"
            t += safe(respuesta, "[Completar información solicitada]") + "\n"

            if adjuntos.strip():
                t += "\nIII. Documentación adjunta:\n" + adjuntos.strip() + "\n"

            if requiere_consent:
                t += "\nSe deja constancia que la presente información se brinda en el marco de las facultades y autorizaciones correspondientes.\n"

            if confidencialidad:
                t += "\nLa presente contestación se emite a los fines del requerimiento indicado, para uso exclusivo del organismo requirente.\n"

            t += "\nSin otro particular, saludo atentamente.\n"
            t += bloque_firma(firmante, matricula, estudio, contacto)

        st.session_state["ultimo_oficio"] = t
        st.session_state["sync_editor_oficio"] = True

        guardar_en_historial(
            tipo="Contestación de Oficio",
            titulo=f"Oficio - {organismo or 'Sin organismo'}",
            contenido=t
        )

    if "ultimo_oficio" in st.session_state:

        if "editor_oficio" not in st.session_state:
            st.session_state["editor_oficio"] = st.session_state["ultimo_oficio"]

        if st.session_state.get("sync_editor_oficio", False):
            st.session_state["editor_oficio"] = st.session_state["ultimo_oficio"]
            st.session_state["sync_editor_oficio"] = False

        st.markdown("### Resultado")

        texto_actual_oficio = st.text_area(
            "Texto generado / editable",
            height=420,
            key="editor_oficio"
        )

        st.session_state["ultimo_oficio"] = texto_actual_oficio

        st.markdown("### Editar con IA")
        instruccion_oficio = st.text_input(
            "Pedile cambios a la IA",
            value=st.session_state.get("instruccion_edicion_oficio", ""),
            placeholder="Ej: hacelo más formal, más breve, agregá un cierre más técnico."
        )

        if st.button("Aplicar cambios con IA", key="editar_oficio_ia"):
            if not instruccion_oficio.strip():
                st.warning("Escribí una instrucción para editar el texto.")
            else:
                texto_editado_oficio = editar_texto_con_ia(texto_actual_oficio, instruccion_oficio)

                if not texto_editado_oficio:
                    st.error("No se encontró OPENAI_API_KEY en Secrets.")
                elif str(texto_editado_oficio).startswith("ERROR_IA:"):
                    st.error(texto_editado_oficio)
                else:
                    st.session_state["ultimo_oficio"] = texto_editado_oficio
                    st.session_state["sync_editor_oficio"] = True

                    guardar_en_historial(
                        tipo="Edición IA - Contestación de Oficio",
                        titulo=f"Edición IA - {organismo or 'Sin organismo'}",
                        contenido=texto_editado_oficio
                    )

                    st.success("Texto actualizado con IA.")
                    st.rerun()

        exportar_word(
            st.session_state["ultimo_oficio"],
            "Contestacion_Oficio_Estudio_Peire"
        )

# =========================================================
# 4) MAILING MODO AGENTE
# =========================================================
elif menu == "Mailing (Modo Agente)":
    st.header("📧 Mailing (Modo Agente)")

    col_a, col_b = st.columns([1, 1])
    with col_a:
        st.button("← Volver al panel principal", on_click=volver_al_dashboard)
    with col_b:
        if st.button("Nuevo mailing", key="reset_mail"):
            limpiar_resultado("ultimo_mail")
            limpiar_resultado("editor_mail")
            limpiar_resultado("sync_editor_mail")
            limpiar_resultado("instruccion_edicion_mail")
            st.rerun()
    
    col1, col2, col3 = st.columns(3)
    with col1:
        tipo_mail = st.selectbox("Tipo de mensaje", ["Actualización de caso", "Pedido de documentación", "Seguimiento", "Cierre / próximos pasos", "Recordatorio"])
    with col2:
        tono = st.selectbox("Tono", ["Cálido y profesional", "Muy formal", "Breve y directo"])
    with col3:
        canal = st.selectbox("Canal", ["Email", "WhatsApp (texto corto)"])

    cliente = st.text_input("Nombre del cliente", placeholder="Ej: María González")
    caso = st.text_input("Caso/Asunto general (ej: alquiler, laboral, daños)", placeholder="Ej: Reclamo laboral")
    estado = st.text_area("Estado actual / contexto (2–6 líneas)", height=100, placeholder="Describí el estado actual del caso.")
    proximo_paso = st.text_area("Próximo paso (qué tiene que pasar ahora)", height=80, placeholder="Ej: aguardar respuesta / presentar documentación.")
    accion_cliente = st.text_input("Acción requerida al cliente (si aplica)", value="", placeholder="Ej: enviar DNI y comprobantes")
    col4, col5 = st.columns(2)
    with col4:
        incluir_disclaimer = st.checkbox("Incluir disclaimer (confidencialidad)", value=True)
    with col5:
        incluir_agenda = st.checkbox("Sugerir coordinación de llamada/reunión", value=False)

    usar_ia_mail = st.checkbox("Usar IA para redactar el mensaje", value=True)
    
    if st.button("Generar Mailing"):

        limpiar_resultado("ultimo_mail")
        
        if usar_ia_mail:
            prompt_sistema = (
                "Sos asistente del Estudio Peire. "
                "Redactás emails y mensajes a clientes en español claro, profesional y útil. "
                "No inventes hechos. "
                "Devolvé solo el texto final."
            )

            prompt_usuario = f"""
Redactá una comunicación para cliente.

Tipo de mensaje: {tipo_mail}
Canal: {canal}
Tono: {tono}
Cliente: {cliente}
Caso: {caso}

Estado actual:
{estado}

Próximo paso:
{proximo_paso}

Acción requerida al cliente:
{accion_cliente}

Sugerir coordinación de llamada/reunión: {incluir_agenda}
Incluir disclaimer de confidencialidad: {incluir_disclaimer}

Estudio: {estudio}
Contacto: {contacto}

Devolvé solo el texto final del mensaje.
"""
            t = generar_texto_con_ia(prompt_sistema, prompt_usuario)

            if not t:
                st.error("No se encontró OPENAI_API_KEY en Secrets.")
                st.stop()

            if str(t).startswith("ERROR_IA:"):
                st.error(t)
                st.stop()

        else:
            nombre_cli = safe(cliente, "[Cliente]")
            caso_txt = safe(caso, "[Caso]")

            if tipo_mail == "Actualización de caso":
                asunto = f"Actualización – {caso_txt}"
            elif tipo_mail == "Pedido de documentación":
                asunto = f"Documentación necesaria – {caso_txt}"
            elif tipo_mail == "Seguimiento":
                asunto = f"Seguimiento – {caso_txt}"
            elif tipo_mail == "Cierre / próximos pasos":
                asunto = f"Próximos pasos – {caso_txt}"
            else:
                asunto = f"Recordatorio – {caso_txt}"

            if canal == "WhatsApp (texto corto)":
                t = f"{nombre_cli}, te escribo desde {estudio}. "
                if tipo_mail == "Pedido de documentación":
                    t += "Necesitamos la siguiente documentación/confirmación para avanzar. "
                t += safe(estado, "Actualización del caso. ").strip() + " "
                if accion_cliente.strip():
                    t += f"¿Podés enviarnos: {accion_cliente.strip()}? "
                if proximo_paso.strip():
                    t += f"Próximo paso: {proximo_paso.strip()} "
                if incluir_agenda:
                    t += "Si te parece, coordinamos una llamada breve. "
                if incluir_disclaimer:
                    t += "Mensaje confidencial."
            else:
                if tono == "Cálido y profesional":
                    saludo = f"Hola {nombre_cli},"
                    cierre = "Un saludo"
                elif tono == "Muy formal":
                    saludo = f"De mi mayor consideración {nombre_cli}:"
                    cierre = "Atentamente"
                else:
                    saludo = f"{nombre_cli},"
                    cierre = "Saludos"

                t = f"Asunto: {asunto}\n\n{saludo}\n\n"
                t += safe(estado, "[Estado actual del caso]") + "\n\n"

                if tipo_mail == "Pedido de documentación" and accion_cliente.strip():
                    t += f"Para poder avanzar, necesitamos que nos envíes: {accion_cliente.strip()}.\n\n"
                elif accion_cliente.strip():
                    t += f"Acción requerida: {accion_cliente.strip()}.\n\n"

                if proximo_paso.strip():
                    t += f"Próximo paso: {proximo_paso.strip()}.\n\n"

                if incluir_agenda:
                    t += "Si estás de acuerdo, coordinamos una llamada/reunión breve para confirmar los próximos pasos.\n\n"

                if incluir_disclaimer:
                    t += "Este mensaje contiene información confidencial. Si no sos el destinatario, por favor informanos y eliminá el contenido.\n\n"

                t += f"{cierre},\n{estudio}\n"
                if contacto.strip():
                    t += f"{contacto}\n"

        st.session_state["ultimo_mail"] = t
        st.session_state["sync_editor_mail"] = True

        guardar_en_historial(
            tipo="Mailing",
            titulo=f"Mailing - {cliente or 'Sin cliente'}",
            contenido=t
        )

    if "ultimo_mail" in st.session_state:

        if "editor_mail" not in st.session_state:
            st.session_state["editor_mail"] = st.session_state["ultimo_mail"]

        if st.session_state.get("sync_editor_mail", False):
            st.session_state["editor_mail"] = st.session_state["ultimo_mail"]
            st.session_state["sync_editor_mail"] = False

        st.markdown("### Resultado")

        texto_actual_mail = st.text_area(
            "Texto generado / editable",
            height=320,
            key="editor_mail"
        )

        st.session_state["ultimo_mail"] = texto_actual_mail

        st.markdown("### Editar con IA")
        instruccion_mail = st.text_input(
            "Pedile cambios a la IA",
            value=st.session_state.get("instruccion_edicion_mail", ""),
            placeholder="Ej: hacelo más breve, más formal, más cercano."
        )

        if st.button("Aplicar cambios con IA", key="editar_mail_ia"):
            if not instruccion_mail.strip():
                st.warning("Escribí una instrucción para editar el texto.")
            else:
                texto_editado_mail = editar_texto_con_ia(texto_actual_mail, instruccion_mail)

                if not texto_editado_mail:
                    st.error("No se encontró OPENAI_API_KEY en Secrets.")
                elif str(texto_editado_mail).startswith("ERROR_IA:"):
                    st.error(texto_editado_mail)
                else:
                    st.session_state["ultimo_mail"] = texto_editado_mail
                    st.session_state["sync_editor_mail"] = True

                    guardar_en_historial(
                        tipo="Edición IA - Mailing",
                        titulo=f"Edición IA - {cliente or 'Sin cliente'}",
                        contenido=texto_editado_mail
                    )

                    st.success("Texto actualizado con IA.")
                    st.rerun()

        exportar_word(
            st.session_state["ultimo_mail"],
            "Mailing_Estudio_Peire"
        )

# =========================================================
# 5) PRESUPUESTO
# =========================================================
elif menu == "Presupuesto":
    st.header("💼 Presupuesto de Honorarios")

    col_a, col_b = st.columns([1, 1])
    with col_a:
        st.button("← Volver al panel principal", on_click=volver_al_dashboard)
    with col_b:
        if st.button("Nuevo presupuesto", key="reset_presupuesto"):
            limpiar_resultado("ultimo_presupuesto")
            limpiar_resultado("editor_presupuesto")
            limpiar_resultado("sync_editor_presupuesto")
            limpiar_resultado("instruccion_edicion_presupuesto")
            st.rerun()
    
    col1, col2 = st.columns(2)
    with col1:
        cliente = st.text_input("Cliente")
        servicio = st.text_input("Servicio")
    with col2:
        fecha = st.text_input("Fecha (dd/mm/aaaa)", value=date.today().strftime("%d/%m/%Y"))
        validez = st.selectbox("Validez del presupuesto", ["7 días", "10 días", "15 días", "30 días"])

    modalidad = st.selectbox("Modalidad", ["Monto fijo", "Por etapas", "Success fee", "Mixto (fijo + success fee)"])
    honorarios = st.text_input("Honorarios / Monto / Porcentaje (según modalidad)")

    cliente = st.text_input("Cliente", placeholder="Ej: María González")
    servicio = st.text_input("Servicio", placeholder="Ej: Sucesión / Carta documento / Reclamo")
    honorarios = st.text_input("Honorarios / Monto / Porcentaje (según modalidad)", placeholder="Ej: $250.000 o 10%")
    alcance = st.text_area("Alcance (qué incluye)", height=110, placeholder="Describí qué incluye el servicio.")
    no_incluye = st.text_area("No incluye (limitaciones)", height=90, placeholder="Describí qué no está incluido.")
    plazos = st.text_area("Plazos estimados", height=80, placeholder="Ej: entre 30 y 60 días.")
    forma_pago = st.text_area("Forma de pago", height=80, placeholder="Ej: 50% al inicio y 50% contra entrega.")
    observaciones = st.text_area("Observaciones (opcional)", height=70, placeholder="Aclaraciones adicionales.")
    col3, col4, col5 = st.columns(3)
    with col3:
        incluir_impuestos = st.checkbox("Aclarar impuestos/retenciones (si aplica)", value=True)
    with col4:
        incluir_gastos = st.checkbox("Aclarar gastos (tasa, diligencias, etc.)", value=True)
    with col5:
        incluir_condiciones = st.checkbox("Condiciones generales (cambios de alcance)", value=True)

    observaciones = st.text_area("Observaciones (opcional)", height=70)

    usar_ia_presupuesto = st.checkbox("Usar IA para redactar el presupuesto", value=True)
    
    if st.button("Generar Presupuesto"):

        limpiar_resultado("ultimo_presupuesto")
        
        if usar_ia_presupuesto:
            prompt_sistema = (
                "Sos asistente del Estudio Peire. "
                "Redactás presupuestos jurídicos profesionales en español claro y formal. "
                "No inventes hechos ni condiciones no dadas. "
                "Devolvé solo el texto final."
            )

            prompt_usuario = f"""
Redactá un presupuesto de honorarios jurídicos.

Cliente: {cliente}
Servicio: {servicio}
Fecha: {fecha}
Validez: {validez}
Modalidad: {modalidad}
Honorarios: {honorarios}

Alcance:
{alcance}

No incluye:
{no_incluye}

Plazos:
{plazos}

Forma de pago:
{forma_pago}

Aclarar impuestos/retenciones: {incluir_impuestos}
Aclarar gastos: {incluir_gastos}
Incluir condiciones generales: {incluir_condiciones}

Observaciones:
{observaciones}

Estudio: {estudio}
Firmante: {firmante}
Matrícula: {matricula}
Contacto: {contacto}

Devolvé solo el texto final del presupuesto.
"""
            t = generar_texto_con_ia(prompt_sistema, prompt_usuario)

            if not t:
                st.error("No se encontró OPENAI_API_KEY en Secrets.")
                st.stop()

            if str(t).startswith("ERROR_IA:"):
                st.error(t)
                st.stop()

            t += bloque_firma(firmante, matricula, estudio, contacto)

        else:
            t = "PRESUPUESTO DE HONORARIOS\n\n"
            t += f"Estudio: {estudio}\n"
            t += f"Fecha: {safe(fecha,'[Fecha]')}\n"
            t += f"Cliente: {safe(cliente,'[Cliente]')}\n"
            t += f"Servicio: {safe(servicio,'[Servicio]')}\n\n"

            t += f"Modalidad: {modalidad}\n"
            t += f"Honorarios: {safe(honorarios,'[Completar]')}\n\n"

            t += "Alcance (incluye):\n" + safe(alcance, "[Detallar alcance]") + "\n\n"
            t += "No incluye:\n" + safe(no_incluye, "[Detallar exclusiones]") + "\n\n"
            t += "Plazos estimados:\n" + safe(plazos, "[Detallar plazos]") + "\n\n"
            t += "Forma de pago:\n" + safe(forma_pago, "[Detallar forma de pago]") + "\n\n"

            if incluir_gastos:
                t += "Gastos:\nLos gastos y erogaciones no se encuentran incluidos salvo indicación expresa.\n\n"
            if incluir_impuestos:
                t += "Impuestos/retenciones:\nLos importes podrán estar sujetos a impuestos y/o retenciones según normativa aplicable.\n\n"
            if incluir_condiciones:
                t += "Condiciones generales:\nEl presente presupuesto se basa en la información provista. Cualquier ampliación del alcance o complejidad no prevista podrá implicar ajustes.\n\n"

            t += f"Validez: {validez}\n"
            if observaciones.strip():
                t += "\nObservaciones:\n" + observaciones.strip() + "\n"

            t += bloque_firma(firmante, matricula, estudio, contacto)

        st.session_state["ultimo_presupuesto"] = t
        st.session_state["sync_editor_presupuesto"] = True

        guardar_en_historial(
            tipo="Presupuesto",
            titulo=f"Presupuesto - {cliente or 'Sin cliente'}",
            contenido=t
        )

    if "ultimo_presupuesto" in st.session_state:

        if "editor_presupuesto" not in st.session_state:
            st.session_state["editor_presupuesto"] = st.session_state["ultimo_presupuesto"]

        if st.session_state.get("sync_editor_presupuesto", False):
            st.session_state["editor_presupuesto"] = st.session_state["ultimo_presupuesto"]
            st.session_state["sync_editor_presupuesto"] = False

        st.markdown("### Resultado")

        texto_actual_presupuesto = st.text_area(
            "Texto generado / editable",
            height=360,
            key="editor_presupuesto"
        )

        st.session_state["ultimo_presupuesto"] = texto_actual_presupuesto

        st.markdown("### Editar con IA")
        instruccion_presupuesto = st.text_input(
            "Pedile cambios a la IA",
            value=st.session_state.get("instruccion_edicion_presupuesto", ""),
            placeholder="Ej: hacelo más formal, agregá condiciones, resumilo."
        )

        if st.button("Aplicar cambios con IA", key="editar_presupuesto_ia"):
            if not instruccion_presupuesto.strip():
                st.warning("Escribí una instrucción para editar el texto.")
            else:
                texto_editado_presupuesto = editar_texto_con_ia(texto_actual_presupuesto, instruccion_presupuesto)

                if not texto_editado_presupuesto:
                    st.error("No se encontró OPENAI_API_KEY en Secrets.")
                elif str(texto_editado_presupuesto).startswith("ERROR_IA:"):
                    st.error(texto_editado_presupuesto)
                else:
                    st.session_state["ultimo_presupuesto"] = texto_editado_presupuesto
                    st.session_state["sync_editor_presupuesto"] = True

                    guardar_en_historial(
                        tipo="Edición IA - Presupuesto",
                        titulo=f"Edición IA - {cliente or 'Sin cliente'}",
                        contenido=texto_editado_presupuesto
                    )

                    st.success("Texto actualizado con IA.")
                    st.rerun()

        exportar_word(
            st.session_state["ultimo_presupuesto"],
            "Presupuesto_Estudio_Peire"
        )

# =========================================================
# ANÁLISIS DE DOCUMENTO
# =========================================================
elif menu == "Análisis de Documento":
    st.button("← Volver al panel principal", on_click=volver_al_dashboard, key="volver_analisis")

    st.header("📂 Análisis de Documento")

    uploaded_file = st.file_uploader(
        "Subir archivo",
        type=["pdf", "docx", "txt", "jpg", "jpeg", "png"],
        key="archivo_analisis"
    )

    tipo_opciones = [
        "Carta Documento recibida",
        "Respuesta a Carta Documento",
        "Oficio recibido",
        "Intimación",
        "Otro"
    ]

    if "remitente_analisis" not in st.session_state:
        st.session_state["remitente_analisis"] = ""

    if "destinatario_analisis" not in st.session_state:
        st.session_state["destinatario_analisis"] = ""

    if "fecha_doc_analisis" not in st.session_state:
        st.session_state["fecha_doc_analisis"] = ""

    if "monto_analisis" not in st.session_state:
        st.session_state["monto_analisis"] = ""

    if "objeto_analisis" not in st.session_state:
        st.session_state["objeto_analisis"] = ""

    if "resumen_analisis_manual" not in st.session_state:
        st.session_state["resumen_analisis_manual"] = ""

    tipo_documento = st.selectbox(
        "Tipo de documento",
        tipo_opciones,
        key="tipo_documento_analisis"
    )

    observaciones = st.text_area(
        "Observaciones / contexto del estudio",
        height=120,
        placeholder="Ej: este documento llegó hoy, corresponde a un reclamo por alquiler, el cliente dice que ya pagó, etc.",
        key="observaciones_analisis"
    )

    contenido_extraido = ""
    datos_detectados = {}

    if "archivo_analisis_procesado" not in st.session_state:
        st.session_state["archivo_analisis_procesado"] = ""

    if "datos_analisis_cargados" not in st.session_state:
        st.session_state["datos_analisis_cargados"] = False
    
    if uploaded_file is not None:
        nombre_archivo_actual = uploaded_file.name

        with st.spinner("Procesando archivo..."):
            contenido_extraido = extraer_texto_archivo(uploaded_file)

        if contenido_extraido.startswith("ERROR_"):
            st.error(contenido_extraido)
            contenido_extraido = ""

        elif contenido_extraido.strip():
            st.success(f"Archivo cargado: {uploaded_file.name}")

            st.text_area(
                "Texto detectado del archivo",
                value=contenido_extraido,
                height=220,
                key="texto_detectado_analisis"
            )

        # Solo detectar datos si el archivo todavía no fue procesado
            if (
                st.session_state["archivo_analisis_procesado"] != nombre_archivo_actual
                or not st.session_state["datos_analisis_cargados"]
            ):
                with st.spinner("Extrayendo datos clave..."):
                    datos_detectados = extraer_datos_clave_con_ia(contenido_extraido)

                if isinstance(datos_detectados, dict) and "error" in datos_detectados:
                    st.error(datos_detectados["error"])
                    datos_detectados = {}

                elif isinstance(datos_detectados, dict):
                    st.subheader("Datos detectados automáticamente")

                    st.markdown(f"""
**Tipo sugerido:** {datos_detectados.get("tipo_documento", "No detectado")}  
**Remitente:** {datos_detectados.get("remitente", "No detectado")}  
**Destinatario:** {datos_detectados.get("destinatario", "No detectado")}  
**Fecha:** {datos_detectados.get("fecha", "No detectado")}  
**Monto:** {datos_detectados.get("monto", "No detectado")}  
**Objeto:** {datos_detectados.get("objeto", "No detectado")}  
**Resumen:** {datos_detectados.get("resumen", "No detectado")}
""")

                st.session_state["remitente_analisis"] = datos_detectados.get("remitente", "")
                st.session_state["destinatario_analisis"] = datos_detectados.get("destinatario", "")
                st.session_state["fecha_doc_analisis"] = datos_detectados.get("fecha", "")
                st.session_state["monto_analisis"] = datos_detectados.get("monto", "")
                st.session_state["objeto_analisis"] = datos_detectados.get("objeto", "")
                st.session_state["resumen_analisis_manual"] = datos_detectados.get("resumen", "")

                tipo_detectado = datos_detectados.get("tipo_documento", "")
                if tipo_detectado:
                    st.info(f"Tipo sugerido por IA: {tipo_detectado}")

                st.session_state["archivo_analisis_procesado"] = nombre_archivo_actual
                st.session_state["datos_analisis_cargados"] = True

                st.rerun()

            else:
                st.warning("No se pudieron estructurar los datos detectados.")
                datos_detectados = {}

        else:
            st.info("Datos ya detectados para este archivo.")

    else:
        st.warning("No se pudo extraer texto del archivo o está vacío.")

    st.subheader("Datos clave del documento")

    remitente = st.text_input(
        "Remitente",
        key="remitente_analisis"
    )

    destinatario = st.text_input(
        "Destinatario",
        key="destinatario_analisis"
    )

    fecha_doc = st.text_input(
        "Fecha del documento",
        key="fecha_doc_analisis"
    )

    monto = st.text_input(
        "Monto (si aplica)",
        key="monto_analisis",
        placeholder="Ej: $450.000"
    )

    objeto = st.text_input(
        "Objeto / tema principal",
        key="objeto_analisis",
        placeholder="Ej: Reclamo por alquiler adeudado"
    )

    resumen = st.text_area(
        "Resumen manual / puntos importantes",
        height=120,
        key="resumen_analisis_manual",
        placeholder="Ej: intiman pago por alquiler, reclaman $450.000, niegan pagos, dan plazo de 48 hs, etc."
    )

    usar_ia_analisis = st.checkbox(
        "Usar IA para analizar el documento",
        value=True,
        key="usar_ia_analisis"
    )

    col_a, col_b = st.columns([1, 1])

    with col_a:
        if st.button("Preparar borrador con IA", key="generar_analisis"):
            limpiar_resultado("ultimo_analisis_documento")
            limpiar_resultado("editor_analisis_documento")
            limpiar_resultado("sync_editor_analisis_documento")

            texto_base = contenido_extraido if contenido_extraido.strip() else resumen

            if usar_ia_analisis and texto_base.strip():
                prompt_sistema = (
                    "Sos asistente jurídico del Estudio Peire. "
                    "Analizás documentos jurídicos en español argentino. "
                    "No inventes hechos ni normas. "
                    "Ordenás la información con claridad profesional. "
                    "Devolvé un análisis interno útil para el estudio."
                )

                prompt_usuario = f"""
Analizá el siguiente documento y prepará un borrador interno para el estudio.

Tipo de documento: {tipo_documento}
Remitente: {remitente}
Destinatario: {destinatario}
Fecha: {fecha_doc}
Monto: {monto}
Objeto: {objeto}

Observaciones del estudio:
{observaciones}

Resumen manual:
{resumen}

Texto del documento:
{texto_base}

Devolvé:
1. Resumen ejecutivo
2. Puntos clave
3. Riesgos o alertas
4. Estrategia sugerida
5. Próximo paso recomendado
"""
                borrador = generar_texto_con_ia(prompt_sistema, prompt_usuario)

                if not borrador:
                    st.error("No se encontró OPENAI_API_KEY en Secrets.")
                    st.stop()

                if str(borrador).startswith("ERROR_IA:"):
                    st.error(borrador)
                    st.stop()
            else:
                borrador = f"""
ANÁLISIS DEL DOCUMENTO

Tipo de documento: {tipo_documento}
Archivo cargado: {uploaded_file.name if uploaded_file else "[Sin archivo]"}
Remitente: {remitente or "[No informado]"}
Destinatario: {destinatario or "[No informado]"}
Fecha: {fecha_doc or "[No informada]"}
Monto: {monto or "[No informado]"}
Objeto: {objeto or "[No informado]"}

Observaciones del estudio:
{observaciones or "[Sin observaciones]"}

Resumen / puntos importantes:
{resumen or "[Sin resumen]"}

Texto extraído del archivo:
{texto_base or "[Sin texto extraído]"}

SUGERENCIA DE PRÓXIMO PASO:
Se recomienda revisar el contenido del documento y utilizar la información arriba consignada para preparar la respuesta correspondiente dentro del módulo "Respuesta Carta Documento" o "Contestación de Oficio", según corresponda.
"""

            st.session_state["analisis_para_respuesta"] = {
                "texto_recibido": texto_base,
                "hechos_reales": observaciones,
                "remitente": remitente,
                "destinatario": destinatario,
                "fecha_doc": fecha_doc,
                "monto": monto,
                "objeto": objeto,
                "tipo_documento": tipo_documento,
                "resumen": resumen,
            }

            st.session_state["ultimo_analisis_documento"] = borrador
            st.session_state["sync_editor_analisis_documento"] = True

            guardar_en_historial(
                tipo="Análisis de Documento",
                titulo=f"Análisis - {uploaded_file.name if uploaded_file else 'Sin archivo'}",
                contenido=borrador
            )

            st.rerun()

    with col_b:
        if st.button("Limpiar filtros", key="reset_analisis"):
            st.session_state["tipo_documento_analisis"] = "Carta Documento recibida"
            st.session_state["remitente_analisis"] = ""
            st.session_state["destinatario_analisis"] = ""
            st.session_state["fecha_doc_analisis"] = ""
            st.session_state["monto_analisis"] = ""
            st.session_state["objeto_analisis"] = ""
            st.session_state["resumen_analisis_manual"] = ""
            st.session_state["observaciones_analisis"] = ""
            st.session_state["archivo_analisis_procesado"] = ""
            st.session_state["datos_analisis_cargados"] = False

            limpiar_resultado("ultimo_analisis_documento")
            limpiar_resultado("editor_analisis_documento")
            limpiar_resultado("sync_editor_analisis_documento")
            limpiar_resultado("instruccion_edicion_analisis_documento")

            st.rerun()

    if "ultimo_analisis_documento" in st.session_state:

        if "editor_analisis_documento" not in st.session_state:
            st.session_state["editor_analisis_documento"] = st.session_state["ultimo_analisis_documento"]

        if st.session_state.get("sync_editor_analisis_documento", False):
            st.session_state["editor_analisis_documento"] = st.session_state["ultimo_analisis_documento"]
            st.session_state["sync_editor_analisis_documento"] = False

        st.markdown("### Resultado del análisis")

        texto_actual_analisis = st.text_area(
            "Análisis generado / editable",
            height=420,
            key="editor_analisis_documento"
        )

        st.session_state["ultimo_analisis_documento"] = texto_actual_analisis

        st.markdown("### Editar análisis con IA")
        instruccion_edicion_analisis = st.text_input(
            "Pedile cambios a la IA",
            value=st.session_state.get("instruccion_edicion_analisis_documento", ""),
            placeholder="Ej: resumilo más, agregá riesgos, proponé una estrategia más concreta."
        )

        if st.button("Aplicar cambios al análisis con IA", key="editar_analisis_ia"):
            if not instruccion_edicion_analisis.strip():
                st.warning("Escribí una instrucción para editar el análisis.")
            else:
                texto_editado_analisis = editar_texto_con_ia(
                    texto_actual_analisis,
                    instruccion_edicion_analisis
                )

                if not texto_editado_analisis:
                    st.error("No se encontró OPENAI_API_KEY en Secrets.")
                elif str(texto_editado_analisis).startswith("ERROR_IA:"):
                    st.error(texto_editado_analisis)
                else:
                    st.session_state["ultimo_analisis_documento"] = texto_editado_analisis
                    st.session_state["sync_editor_analisis_documento"] = True

                    guardar_en_historial(
                        tipo="Edición IA - Análisis de Documento",
                        titulo=f"Edición IA - {uploaded_file.name if uploaded_file else 'Sin archivo'}",
                        contenido=texto_editado_analisis
                    )

                    st.success("Análisis actualizado con IA.")
                    st.rerun()

        exportar_word(
            st.session_state["ultimo_analisis_documento"],
            "Analisis_Documento_Estudio_Peire"
        ) 

# =========================================================
# 7) HISTORIAL
# =========================================================
elif menu == "Historial":
    st.button("← Volver al panel principal", on_click=volver_al_dashboard)
    st.header("🕘 Historial de Documentos")
    
    historial = st.session_state.get("historial_documentos", [])

    if not historial:
        st.info("Todavía no se generaron documentos en esta sesión.")
    else:
        st.write(f"Se encontraron {len(historial)} documento(s) generados.")

        for i, item in enumerate(historial):
            with st.expander(f"{item['fecha']} | {item['tipo']} | {item['titulo']}"):
                st.text_area(
                    f"Contenido {i+1}",
                    item["contenido"],
                    height=250,
                    key=f"historial_{i}"
                )
                exportar_word(item["contenido"], item["titulo"].replace(" ", "_"))

        if st.button("Borrar historial"):
            st.session_state["historial_documentos"] = []
            st.success("Historial borrado.")
            st.rerun()

st.divider()

if st.button("🧹 Limpiar filtros"):

    claves = list(st.session_state.keys())

    for k in claves:
        if (
            "editor" in k
            or "ultimo" in k
            or "sync" in k
            or "analisis" in k
            or "archivo" in k
        ):
            del st.session_state[k]

    st.rerun()

# =========================================================
# 8) BIBLIOTECA OFICIAL DE PROMPTS
# =========================================================
elif menu == "Biblioteca Oficial de Prompts":
    st.button("⬅ Volver al Dashboard", on_click=volver_al_dashboard)
    st.header("📚 Biblioteca Oficial de Prompts – Estudio Peire")
    
    st.subheader("Prompt maestro (pegar al inicio)")
    st.code(
"""Sos asistente del Estudio Peire. Objetivo: redactar BORRADORES y ordenar información.
Reglas:
- No inventes jurisprudencia, normas ni citas.
- Si faltan datos, preguntá de forma concreta.
- Redactá en español, tono profesional y claro.
- Entregá siempre: (a) texto final editable, (b) checklist de datos a verificar, (c) riesgos/puntos sensibles.
- Incluir al final: "Revisar y adecuar por profesional antes de enviar/presentar". """
    )

    st.subheader("Carta Documento")
    st.code(
"""Necesito un borrador de CARTA DOCUMENTO (Argentina), estilo Estudio Peire.
Datos:
- Tipo:
- Remitente + domicilio:
- Destinatario + domicilio:
- Hechos (cronología breve):
- Monto (si aplica):
- Plazo intimado:
- Pedido concreto:
- Documentación/pruebas:
- Tono: neutral/firme/muy firme
Entregá:
1) Texto listo
2) Versión alternativa más firme
3) Checklist y riesgos."""
    )

    st.subheader("Respuesta a Carta Documento")
    st.code(
"""Redactá RESPUESTA a carta documento (Argentina), estilo Estudio Peire.
Texto recibido:
[...]
Hechos reales del cliente:
[...]
Objetivo: negar / aceptar parcial / proponer acuerdo / intimar
Tono: neutral/firme/muy firme
Entregá texto + puntos sensibles + preguntas faltantes."""
    )

    st.subheader("Contestación de Oficio")
    st.code(
"""Borrador de CONTESTACIÓN DE OFICIO.
Organismo/Juzgado:
Expte/Carátula:
Pedido del oficio:
Datos a informar:
Adjuntos:
¿Confidencialidad/consentimiento?:
Entregá texto listo + campos a completar + advertencias."""
    )

    st.subheader("Mailing modo agente")
    st.code(
"""Actuá como agente de atención al cliente del Estudio Peire.
Objetivo: actualización / pedir docs / seguimiento / cierre
Cliente:
Caso:
Estado actual:
Acción requerida:
Próximo paso:
Tono: cálido / formal / breve
Entregá: Email + WhatsApp corto + versión formal."""
    )

    st.subheader("Presupuesto")
    st.code(
"""Generá un presupuesto estilo Estudio Peire.
Cliente:
Servicio:
Modalidad:
Honorarios:
Incluye:
No incluye:
Plazos:
Forma de pago:
Validez:
Entregá presupuesto listo + variables que cambian costo + texto breve para WhatsApp."""
    )
